# BIFS617 ORF Finder Skeleton
# Team members:Alexis Yanes, Bash Harb, Dennis Kim
# Date:07/31/2025
# Fill in your code where prompted.


# Constants for start and stop codons — these define where an ORF can begin and end
START_CODON = 'ATG'
STOP_CODONS = {'TAA', 'TAG', 'TGA'}

def load_fasta(filepath):
    # Team Member Name:Alexis Yanes
    # TODO: Parse a multi-line FASTA file, return dictionary {header: sequence}
    sequence = {}
    header_row = None #default none if no header is present
    with open (filepath, 'r') as f: #Path file and read
        for row in f:
            row = row.strip() # removes any new characters or spaces
            
            if row.startswith('>'): #fasta format headers begin with '>'
                header_row = row [1:] #slices to create new string from index 1 to the end (removes >)
                sequence [header_row] = ''# empty header is created in the sequence dictionary. Header will be stored as it's corresponding sequence
            else: #if line doesn't start with '>', it's sequence data
                sequence [header_row] += row.upper() #appends header to the sequence and converts to upper-case
    return sequence

# Team Member Name:Basheer Harb
# TODO: Return reverse complement of sequence (optional: use Bio.Seq)
# Reverse complement using Bio.Seq module
def reverse_complement(seq):
    return str(Seq(seq).reverse_complement())


def find_orfs(header, sequence, min_len, strand='+'):
    # Team Member Name: [Basheer Harb]
    # This function finds all valid Open Reading Frames (ORFs) in a DNA sequence
    # It checks all 3 reading frames on one strand (either forward or reverse)

    start_codon = 'ATG'  # Start codon in DNA
    stop_codons = {'TAA', 'TAG', 'TGA'}  # Valid stop codons

    sequence = sequence.upper()  # Convert the sequence to uppercase 
    orfs = []  # List to store all ORFs found
    seq_len = len(sequence) # Store the total length of the DNA sequence

    # Loop through each of the 3 reading frames: 0, 1, and 2
    for frame in range(3):
        pos = frame  # Start at position 0, 1, or 2

        # Iterate through the sequence in steps of 3 (codon length) until the end
        while pos <= seq_len - 3:
            codon = sequence[pos:pos+3]  # Read a codon (3 bases)

            if codon == start_codon: # Check if it's a start codon (ATG)
                end = pos + 3  # Start looking after ATG
                
                # Search ahead in steps of 3 for a stop codon
                while end <= seq_len - 3:
                    stop_codon = sequence[end:end+3]  # Get the next codon
                    if stop_codon in stop_codons: # Found a valid stop codon
                        orf_seq = sequence[pos:end+3]  # Extract the full ORF including stop codon
                        
                        
                        if len(orf_seq) >= min_len * 3: # Check if ORF is at least the minimum required length
                            # Calculate position based on strand direction
                            position = pos if strand == '+' else seq_len - end - 3
                            # Store ORF information as a tuple: (header, frame, start position, sequence, strand)
                            orfs.append((header, frame+1, position, orf_seq, strand))
                        break  # Stop after first valid stop
                    
                    end += 3 # Move to the next codon

                pos += 3  # Move past the entire ORF to avoid overlapping start codons
            else:
                pos += 3  # Not a start codon, go to the next codon
    # Return the list of all detected ORFs
    return orfs


def format_orf_output(header, frame, position, seq, strand):
    # Team Member Name: [Alexis+Basheer+Dennis]
    # Returns a FASTA-formatted header and sequence with codons separated by space
    # Split the ORF sequence into codons (triplets of 3 bases) and join them with spaces
    codon_seq = ' '.join([seq[i:i+3] for i in range(0, len(seq), 3)])
    
    # Create the FASTA-style header with metadata:
    fasta_header = f">{header} | FRAME = {frame} | POS = {position} | LEN = {len(seq)} | {strand}"
    
    # Return the full FASTA-formatted string: header followed by the codon-separated sequence
    return f"{fasta_header}\n{codon_seq}\n"

def main():
    # Team Member Name:(Alexis & Bash & Dennis)
    # TODO: Implement user input, sequence processing, and ORF printing and save the file
    # Prompt the user to enter the FASTA file name (with a suggested example)
    fasta_file = input("Enter the name of the FASTA file (e.g., orfs_test_input.fasta): ").strip()
    if not fasta_file:  # If user presses Enter without input
        print("No file entered. Exiting.")
        return
    if not os.path.exists(fasta_file):  # Check if the file actually exists
        print(f"File '{fasta_file}' not found. Please check the name and try again.")
        return

    # Prompt the user for the minimum ORF length in codons (optional)
    user_m = input("Enter minimum ORF length in codons (press ENTER to use default = 5): ").strip()
    min_len = int(user_m) if user_m.isdigit() else 5  # Use user input or default
   
    sequences = load_fasta(fasta_file)
    all_orfs = []

    # Process each sequence
    for header, seq in sequences.items():
        # Forward strand . Call function find_orfs in forward strand
        all_orfs.extend(find_orfs(header, seq, min_len, '+'))

        # Reverse strand . Call function reverse_complement(seq) 
        rev_seq = reverse_complement(seq)
        # Reverse strand . Call function find_orfs in reverse strand
        all_orfs.extend(find_orfs(header, rev_seq, min_len, '-'))

    # Output directories : one for orfs output,one for visualization output, and one for combined output  .
    os.makedirs("output/orfs", exist_ok=True)
    os.makedirs("output/visualization", exist_ok=True)

    # Write results to output file
  
    with open("output/orfs/orf_output.fasta", "w") as out_file:
        for orf in all_orfs:
            out_file.write(format_orf_output(*orf))

    # Print results to console : Call format_orf_output function.
    for orf in all_orfs:
        print(format_orf_output(*orf))
    
# Write summary counts to the output file as well
    with open("output/orfs/orf_output.fasta", "a") as out_file:
        out_file.write("\n# ====== ORF Summary ======\n\n")   
        printed_headers = set()
        for orf in all_orfs:
            header, _, _, _, strand = orf
            if header not in printed_headers:
                printed_headers.add(header)
                forward_count = sum(1 for o in all_orfs if o[0] == header and o[4] == '+')
                reverse_count = sum(1 for o in all_orfs if o[0] == header and o[4] == '-')
                out_file.write(f"> {header}\n{forward_count} ORFs in forward, {reverse_count} ORFs in reverse\n\n")
                print(f"> {header}\n{forward_count} ORFs in forward, {reverse_count} ORFs in reverse\n")
            
            
def visualization():
    # Team Member Name: Dennis Kim
    #Chart the lengths of the sequences that where extracted into "orf_output.fasta"
    #Generate char in native python and matplotlib if user prompts
    currentdirectory=os.getcwd() # get file path
    outputfilepath=currentdirectory+"\\output\\orfs\\orf_output.fasta" #get file path to output file
    modification_time=os.path.getmtime(outputfilepath)#get time output file was created
    current_time=time.time()#current time
    fileage=current_time-modification_time#age of the output file
    if fileage<1:#If the output file was not recetnly generated stop visualization process to prevent visualization of older output file
        orf_dict={} #header(key) and sequence length(value) dictionary
        with open (outputfilepath,"r") as file:#read output file generated from main()
            content=file.read()
        expression = r"(?P<XValue>.+) LEN = (?P<YValue>\d+)" #expression to extract key and value
        for match in re.finditer(expression,content):
            key = match.group("XValue")
            value = match.group("YValue")
            orf_dict[key] = value
        K=max(len(key) for key in orf_dict.keys()) + 10 #minimum length of key (this is to ensure all bars are aligned)
        padding=" "#will add spaces so that bars are aglinged
        orf_dictmod = {key.ljust(K, padding): value for key, value in orf_dict.items()}#adjust key so that bar chart is algined
##############BUILDING OF BAR CHART##############
        original=sys.stdout# export results of program to file step 1
        with open (currentdirectory+"\\output\\visualization\\orf_visualization.txt","w") as f:
            sys.stdout = f# export results of program to file step 2
            max_length = int(max(orf_dictmod.values())) #max length of sequence
            print("#" * (max_length + 50))
            print("SEQUENCE LENGTHS CHART:")#print table label
            print("#" * (max_length + 50))
            print("-"*5+"ORF Header"+"-"*5+" "*(K-19)+"-"*4+"Sequence Length"+"-"*4)

            for item, length in orf_dictmod.items():#print dictionary information to create bar chart
                bar = '=' * int(length)
                print(f"{item.ljust(15)} | {bar} ({length})")

            print("#" * (max_length + 50))

            sys.stdout=original # export results of program to file step 3
        with open(currentdirectory+"\\output\\visualization\\orf_visualization.txt", 'r') as file: #open file created to be displayed on the console
            for line in file:
                print(line)
        while True:#will ask user if they want to open a bar graph using matplotlib
            askforgraph=input("would you like a graph made in matplot?(Y/N):").lower()
            import numpy as np
            if askforgraph=="y":#generate chart based on parameters below
                orf_names = list(orf_dict.keys())
                orf_lengths = [int(length) for length in list(orf_dict.values())]
                max_length_plot = max(orf_lengths)
                plt.yticks(np.arange(0, max_length_plot + 1, 1))
                plt.bar(range(len(orf_names)), orf_lengths, tick_label=orf_names)
                plt.xlabel("ORF Header")
                plt.ylabel("Sequence Length")
                plt.title("ORF Sequence Lengths")
                plt.xticks(rotation=45, ha='right') # Optional: Rotate x-axis labels for better readability
                plt.tight_layout()
                plt.savefig('output/visualization/Matplotgraph.png')
                plt.show() # Display chart
                break
            elif askforgraph=="n":
                print("")
                break
            else:
                print ("Invalid input, Please enter 'Y' or 'N'")
    else:
        print ("")
        
def combinevisual():
    # Team Member Name: Dennis Kim
    #combine output file and visualization file into single file
    currentdirectory=os.getcwd()#get current directory
    visualization=currentdirectory+"\\output\\visualization" #output file from visualization() location
    source_folder=[visualization]
    document = Document()  # Create a new Word document
    for folder in source_folder: #look for txt file, fasta file, and png file
        for filepath in glob.glob(os.path.join(folder, "*.txt")) + \
                        glob.glob(os.path.join(folder, "*.png")):
            
            if filepath.endswith(".txt") or filepath.endswith(".fasta"):
                with open(filepath, "r") as infile:
                    document.add_paragraph(infile.read())  # Add text content as a paragraph
            elif filepath.endswith(".png"):
                document.add_picture(filepath, width=Inches(6),height=Inches(6))  # Add the image, setting width as needed
    
    document.save(os.path.join(currentdirectory, "output\\visualization", "ORF_Visualization.docx")) # Save the document            

if __name__ == "__main__":
    try:
        from Bio import SeqIO

    #for record in SeqIO.parse("orfs_test_input.fasta", "fasta"):
    #print(record.id)
    #print(record.seq.reverse_complement())

        import os
        import re
        import sys
        import matplotlib.pyplot as plt
        from Bio.Seq import Seq  # Importing Seq object from Biopython
        import time
        import glob
        from docx import Document
        from docx.shared import Inches
        main()
        visualization()
        combinevisual()
        input("Press ENTER to exit")
    except ModuleNotFoundError:
        print ("BioPython, Matplotlip, or Docx did not import or module is missing, pleae install module for program to work")
        print ("To Install Bio Python open CMD and input pip install biopython")
        print ("To Install matplotlip open CMD and input pip install matplotlib")
        print ("To Install docx open CMD and input pip install python-docx")
        input("Press ENTER to exit")

